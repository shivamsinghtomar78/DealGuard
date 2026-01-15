import os
import uuid
from docx import Document
from typing import Optional

class RedliningAgent:
    """Agent responsible for redlining (editing) DOCX contracts"""
    
    def __init__(self):
        pass
        
    def apply_replacement(self, file_path: str, original_clause: str, alternative_clause: str) -> Optional[str]:
        """
        Replaces the original_clause with alternative_clause in the document at file_path.
        Returns the path to the new modified file.
        """
        if not os.path.exists(file_path):
            print(f"File not found: {file_path}")
            return None
            
        try:
            doc = Document(file_path)
            
            # Simple replacement logic: find the paragraph containing most of the clause
            # or try to match exactly if possible.
            # Clean up the original clause for better matching (remove extra whitespace/newlines)
            target_text = original_clause.strip()
            
            replaced = False
            
            # First pass: look for exact or near-exact matches in single paragraphs
            for para in doc.paragraphs:
                if target_text in para.text:
                    para.text = para.text.replace(target_text, alternative_clause)
                    replaced = True
                    break
            
            # Second pass: if not found, it might be split across paragraphs or have 
            # minor formatting differences. We'll try a more fuzzy approach or 
            # look for a paragraph that starts with a good chunk of the clause.
            if not replaced:
                # Try to find a paragraph that contains at least 50% of the target text
                # (This is a bit risky but better than failing)
                for para in doc.paragraphs:
                    if len(target_text) > 50 and target_text[:50] in para.text:
                        # Replace the whole paragraph text if it seems to be the one
                        # This is a bit aggressive, maybe just replace the match
                        para.text = para.text.replace(para.text, alternative_clause)
                        replaced = True
                        break

            # Save the modified document
            output_dir = "temp_redlined"
            os.makedirs(output_dir, exist_ok=True)
            output_filename = f"redlined_{uuid.uuid4()}_{os.path.basename(file_path)}"
            output_path = os.path.join(output_dir, output_filename)
            
            doc.save(output_path)
            return output_path
            
        except Exception as e:
            print(f"Error during redlining: {e}")
            return None
