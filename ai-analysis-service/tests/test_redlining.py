from docx import Document
import os
import sys

# Add the app directory to sys.path to import RedliningAgent
sys.path.append(os.path.abspath(os.path.join(os.path.dirname(__file__), '..')))

from app.agents.redlining_agent import RedliningAgent

def create_test_docx(path):
    doc = Document()
    doc.add_heading('Test Contract', 0)
    doc.add_paragraph('This is a test contract.')
    doc.add_paragraph('Section 1: Termination')
    doc.add_paragraph('Either party may terminate this agreement with 30 days notice.')
    doc.add_paragraph('Section 2: Liability')
    doc.add_paragraph('The vendor shall be liable for all damages without limit.')
    doc.save(path)
    print(f"Created test document at {path}")

def test_redlining():
    test_file = "test_contract.docx"
    create_test_docx(test_file)
    
    agent = RedliningAgent()
    
    original_clause = "The vendor shall be liable for all damages without limit."
    alternative_clause = "The vendor's total liability shall be limited to the fees paid under this agreement."
    
    redlined_path = agent.apply_replacement(test_file, original_clause, alternative_clause)
    
    if redlined_path and os.path.exists(redlined_path):
        print(f"Redlined document saved at {redlined_path}")
        
        # Verify the content
        doc = Document(redlined_path)
        found = False
        for para in doc.paragraphs:
            if alternative_clause in para.text:
                found = True
                print(f"Found alternative clause: {para.text}")
                break
        
        if found:
            print("SUCCESS: Clause was replaced correctly.")
        else:
            print("FAILURE: Alternative clause not found in redlined document.")
            
        # Cleanup
        # os.remove(test_file)
        # os.remove(redlined_path)
    else:
        print("FAILURE: Redlined document was not created.")

if __name__ == "__main__":
    test_redlining()
